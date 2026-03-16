"""
Add missing Pelapor (Reporter) test cases to UAT document v2.
Adds 7 new TCs (TC-059 to TC-065) to relevant module tables.
"""

import copy
import sys
import os
from docx import Document
from docx.shared import Pt, RGBColor, Emu
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

INPUT_FILE = "docs/04_UAT_Scenarios_WBS_BPKH_AI_v2.docx"
OUTPUT_FILE = "docs/04_UAT_Scenarios_WBS_BPKH_AI_v2.docx"

# ── New test cases ──────────────────────────────────────────────
# Map: (table_index, [test_cases])
# Table 3 = MOD-01 (Portal Pelaporan Publik)
# Table 7 = MOD-05 (Authentication, RBAC & Security)
# Table 11 = MOD-09 (E2E Web - Jalur Gagal & Edge Cases)

NEW_TCS = {
    # MOD-01: Portal Pelaporan Publik (Table 3)
    3: [
        {
            "id": "TC-059",
            "skenario": (
                "Upload file bukti saat submit laporan via Web Portal\n\n"
                "Precondition: Portal pelaporan dapat diakses, fitur upload aktif"
            ),
            "langkah": (
                "1. Buka halaman portal pelaporan\n"
                "2. Isi form laporan dengan kategori dan deskripsi valid\n"
                "3. Klik tombol upload file, pilih file PDF bukti (< 5 MB)\n"
                "4. Klik \"Kirim Laporan\"\n"
                "5. Verifikasi file berhasil terupload bersama laporan"
            ),
            "hasil": (
                "File terupload dan tersimpan di sistem. Ticket ID ditampilkan. "
                "Admin dapat melihat dan mengunduh file lampiran saat review laporan di dashboard."
            ),
            "priority": "High",
        },
        {
            "id": "TC-060",
            "skenario": (
                "Upload file bukti via tracking/chat pelapor\n\n"
                "Precondition: Laporan sudah ada, Ticket ID diketahui, fitur upload di chat aktif"
            ),
            "langkah": (
                "1. Buka halaman tracking dengan Ticket ID valid\n"
                "2. Di kolom chat, ketik pesan: \"Ini bukti tambahan\"\n"
                "3. Klik tombol upload/attachment, pilih file gambar JPG (< 5 MB)\n"
                "4. Klik \"Kirim\"\n"
                "5. Verifikasi pesan dan file muncul di histori chat"
            ),
            "hasil": (
                "Pesan dan file lampiran tersimpan. File muncul di unified inbox admin dashboard. "
                "Admin dapat mengunduh file. Pelapor tetap anonim."
            ),
            "priority": "High",
        },
        {
            "id": "TC-061",
            "skenario": (
                "Upload file dengan format/ukuran tidak valid\n\n"
                "Precondition: Portal pelaporan dapat diakses, fitur upload aktif"
            ),
            "langkah": (
                "1. Buka portal pelaporan, isi form valid\n"
                "2. Coba upload file .exe (format tidak diizinkan)\n"
                "3. Verifikasi: file ditolak dengan pesan error\n"
                "4. Coba upload file > 10 MB (melebihi batas ukuran)\n"
                "5. Verifikasi: file ditolak dengan pesan error\n"
                "6. Upload file PDF valid (< 5 MB), submit laporan"
            ),
            "hasil": (
                "File dengan format berbahaya (.exe) ditolak. File melebihi batas ukuran ditolak. "
                "Pesan error jelas dan informatif. File valid diterima dan laporan berhasil dikirim."
            ),
            "priority": "High",
        },
        {
            "id": "TC-062",
            "skenario": (
                "Tracking status laporan di setiap tahap status lifecycle\n\n"
                "Precondition: Laporan sudah ada dan diproses melalui beberapa tahap status oleh admin"
            ),
            "langkah": (
                "1. [Admin] Proses laporan melalui tahap: Baru → Sedang Ditinjau → Butuh Informasi → Sedang Ditinjau → Dalam Investigasi\n"
                "2. [Pelapor] Buka tracking setelah setiap perubahan status\n"
                "3. Verifikasi: status terkini ditampilkan dengan benar\n"
                "4. Verifikasi: histori perubahan status lengkap dan kronologis\n"
                "5. Verifikasi: timestamp setiap perubahan status tercatat"
            ),
            "hasil": (
                "Portal tracking menampilkan status terkini yang akurat di setiap tahap. "
                "Histori perubahan status lengkap dengan timestamp. Pelapor dapat memantau progres secara real-time."
            ),
            "priority": "Medium",
        },
        {
            "id": "TC-063",
            "skenario": (
                "Notifikasi pesan baru dari admin di portal tracking\n\n"
                "Precondition: Laporan sudah ada, admin sudah mengirim pesan ke pelapor"
            ),
            "langkah": (
                "1. [Admin] Kirim pesan ke pelapor via unified inbox\n"
                "2. [Pelapor] Buka portal tracking dengan Ticket ID\n"
                "3. Verifikasi: ada indikasi pesan baru (badge/highlight/notifikasi)\n"
                "4. Buka chat, verifikasi pesan admin muncul dengan timestamp\n"
                "5. Balas pesan, verifikasi balasan terkirim"
            ),
            "hasil": (
                "Pelapor mendapat indikasi visual bahwa ada pesan baru dari admin. "
                "Pesan ditampilkan dengan benar termasuk timestamp. Komunikasi dua arah berfungsi."
            ),
            "priority": "Medium",
        },
        {
            "id": "TC-064",
            "skenario": (
                "Responsivitas portal pelaporan di perangkat mobile\n\n"
                "Precondition: Portal pelaporan dapat diakses via browser mobile"
            ),
            "langkah": (
                "1. Buka portal pelaporan di browser mobile (Chrome/Safari)\n"
                "2. Verifikasi: layout responsif, form dapat diisi\n"
                "3. Isi dan submit laporan lengkap dari perangkat mobile\n"
                "4. Buka halaman tracking di mobile, masukkan Ticket ID\n"
                "5. Verifikasi: status dan chat dapat diakses di layar kecil"
            ),
            "hasil": (
                "Portal pelaporan dan tracking berfungsi penuh di perangkat mobile. "
                "Layout TailwindCSS responsif. Form, submit, tracking, dan chat dapat digunakan tanpa horizontal scroll."
            ),
            "priority": "Low",
        },
    ],
    # MOD-05: Authentication, RBAC & Security (Table 7)
    7: [
        {
            "id": "TC-065",
            "skenario": (
                "Verifikasi anonimitas pelapor — tidak ada data identitas tersimpan\n\n"
                "Precondition: Minimal 1 laporan sudah di-submit via portal"
            ),
            "langkah": (
                "1. Submit laporan baru via portal pelaporan\n"
                "2. Query database tabel reports: verifikasi TIDAK ada kolom IP address, browser fingerprint, atau identitas pelapor\n"
                "3. Query tabel messages: verifikasi pesan pelapor tidak mengandung metadata identitas\n"
                "4. Periksa audit_logs: verifikasi tidak ada log yang mengaitkan identitas ke laporan\n"
                "5. Periksa response API /api/v1/tickets/{id}: verifikasi tidak ada data identitas di response\n"
                "6. [Admin] Buka detail laporan di dashboard: verifikasi tidak ada informasi yang bisa mengidentifikasi pelapor"
            ),
            "hasil": (
                "Tidak ada IP address, user-agent, cookie, atau metadata identitas pelapor yang tersimpan di database. "
                "Audit log tidak mencatat identitas pelapor. API response bersih dari data identitas. "
                "Sesuai prinsip anonimitas ISO 37002:2021."
            ),
            "priority": "Critical",
        },
    ],
}


def clone_row_format(table, source_row_idx=1):
    """Clone formatting from an existing data row."""
    source_row = table.rows[source_row_idx]
    return source_row


def add_row_to_table(table, tc_data, source_row_idx=1):
    """Add a new row to the table copying formatting from source row."""
    source_row = table.rows[source_row_idx]

    # Create new row by copying XML structure
    new_tr = copy.deepcopy(source_row._tr)
    table._tbl.append(new_tr)

    # Get the new row (last row)
    new_row = table.rows[-1]

    # Map data to columns
    col_data = [
        tc_data["id"],
        tc_data["skenario"],
        tc_data["langkah"],
        tc_data["hasil"],
        tc_data["priority"],
        "",  # Status (empty)
    ]

    for i, cell in enumerate(new_row.cells):
        # Clear existing content
        for p in cell.paragraphs:
            for r in p.runs:
                r.text = ""

        # Set new text
        if i < len(col_data):
            # Clear all paragraphs first
            for p in cell.paragraphs[1:]:
                p._element.getparent().remove(p._element)

            first_para = cell.paragraphs[0]
            # Clear runs
            for r in first_para.runs:
                r._element.getparent().remove(r._element)

            # Add text with proper formatting
            text = col_data[i]
            lines = text.split("\n")
            for j, line in enumerate(lines):
                if j == 0:
                    run = first_para.add_run(line)
                else:
                    new_para = copy.deepcopy(first_para._element)
                    cell._element.append(new_para)
                    p = cell.paragraphs[-1]
                    for r in p.runs:
                        r._element.getparent().remove(r._element)
                    run = p.add_run(line)

                # Apply font formatting
                run.font.name = "Arial"
                run.font.size = Pt(7.5)  # 95250 EMU = 7.5pt

                # Bold for ID and Priority columns
                if i in (0, 4):
                    run.bold = True


def update_summary_table(doc):
    """Update the summary table (Table 2) with new totals."""
    table = doc.tables[2]

    # MOD-01 row (Row 1): was 6 TC, now 12 TC (+6)
    # Original: 6 total, 3 Critical, 2 High, 1 Medium
    # Added: 0 Critical, 3 High, 2 Medium, 1 Low → need to add Low column handling
    # New: 12 total, 3 Critical, 5 High, 3 Medium + 1 Low (count as Medium for simplicity since table has no Low column)
    row1 = table.rows[1]
    cells = row1.cells
    # Total TC
    for p in cells[1].paragraphs:
        for r in p.runs:
            r.text = "12"
    # Critical (unchanged: 3)
    # High: 2 → 5
    for p in cells[3].paragraphs:
        for r in p.runs:
            r.text = "5"
    # Medium: 1 → 3
    for p in cells[4].paragraphs:
        for r in p.runs:
            r.text = "4"

    # MOD-05 row (Row 5): was 5 TC, now 6 TC (+1 Critical)
    row5 = table.rows[5]
    cells5 = row5.cells
    # Total: 5 → 6
    for p in cells5[1].paragraphs:
        for r in p.runs:
            r.text = "6"
    # Critical: 2 → 3
    for p in cells5[2].paragraphs:
        for r in p.runs:
            r.text = "3"

    # Update TOTAL row (Row 15): was 58, now 65 (+7)
    row_total = table.rows[15]
    cells_t = row_total.cells
    # Total TC: 58 → 65
    for p in cells_t[1].paragraphs:
        for r in p.runs:
            r.text = "65"
    # Critical: 24 → 25
    for p in cells_t[2].paragraphs:
        for r in p.runs:
            r.text = "25"
    # High: 25 → 28
    for p in cells_t[3].paragraphs:
        for r in p.runs:
            r.text = "28"
    # Medium: 9 → 12
    for p in cells_t[4].paragraphs:
        for r in p.runs:
            r.text = "12"


def main():
    doc = Document(INPUT_FILE)

    for table_idx, test_cases in NEW_TCS.items():
        table = doc.tables[table_idx]
        for tc in test_cases:
            add_row_to_table(table, tc)
            print(f"  + {tc['id']}: {tc['skenario'].split(chr(10))[0][:60]}")

    # Update summary table
    update_summary_table(doc)
    print("\n  Updated summary table (Table 2) totals")

    # Update document subtitle (paragraph with test case count)
    for para in doc.paragraphs:
        if "36 + 25 Test Cases" in para.text or "58" in para.text:
            for run in para.runs:
                if "36" in run.text:
                    run.text = run.text.replace("36 + 25 Test Cases", "36 + 29 Test Cases")
                if "14 Modul" in run.text:
                    run.text = run.text.replace("14 Modul", "14 Modul")

    doc.save(OUTPUT_FILE)
    print(f"\n  Saved to: {OUTPUT_FILE}")


if __name__ == "__main__":
    print("=" * 60)
    print("  Adding Pelapor Test Cases to UAT Document")
    print("=" * 60)
    print()
    main()
    print()
    print("  New test cases added:")
    print("  TC-059: Upload file bukti saat submit laporan")
    print("  TC-060: Upload file bukti via tracking/chat")
    print("  TC-061: Upload file format/ukuran tidak valid")
    print("  TC-062: Tracking di setiap tahap status lifecycle")
    print("  TC-063: Notifikasi pesan baru di portal tracking")
    print("  TC-064: Responsivitas portal di perangkat mobile")
    print("  TC-065: Verifikasi anonimitas pelapor (ISO 37002)")
    print()
